# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
d = Document(r"c:\EquipmentAccounting\article_work\Odintsov.docx")
print("Paragraphs:", len(d.paragraphs))
for i, p in enumerate(d.paragraphs):
    print(f"{i:3d} | {p.style.name!r:40s} | {p.text[:120]!r}")
